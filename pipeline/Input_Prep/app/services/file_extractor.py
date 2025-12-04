"""
File text extraction service.

Extracts text from various file formats (TXT, MD, PDF, DOCX) and chunks it
for efficient processing.
"""

import os
from pathlib import Path
from typing import List, Optional, Dict, Any
from app.config import settings
from app.utils.hmac_utils import hash_file_sha256
from app.utils.logger import get_logger
from app.models.schemas import FileChunk, FileInfo

logger = get_logger(__name__)

# Try to import PDF and DOCX libraries with graceful fallback
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available. PDF extraction will be disabled.")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX extraction will be disabled.")


def check_library_availability() -> Dict[str, bool]:
    """
    Check which file extraction libraries are available.
    
    Returns:
        Dictionary mapping file types to availability
    """
    return {
        "txt": True,
        "md": True,
        "pdf": PYMUPDF_AVAILABLE,
        "docx": PYTHON_DOCX_AVAILABLE,
    }


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to the TXT file
    
    Returns:
        Extracted text content
    
    Raises:
        IOError: If file cannot be read
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def extract_text_from_md(file_path: str) -> str:
    """
    Extract text from Markdown file.
    
    Markdown files are plain text, so this is identical to TXT extraction.
    
    Args:
        file_path: Path to the MD file
    
    Returns:
        Extracted text content
    """
    return extract_text_from_txt(file_path)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        Extracted text content from all pages
    
    Raises:
        ImportError: If PyMuPDF is not available
        Exception: If PDF cannot be read or is corrupted
    """
    if not PYMUPDF_AVAILABLE:
        raise ImportError("PyMuPDF is not installed. Install with: pip install PyMuPDF")
    
    text_parts = []
    
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                if page_text.strip():  # Only add non-empty pages
                    text_parts.append(page_text)
                    logger.debug(f"Extracted {len(page_text)} chars from page {page_num}")
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise
    
    return '\n'.join(text_parts)


def extract_images_from_pdf(file_path: str, output_dir: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Extract embedded images from PDF file.
    
    Args:
        file_path: Path to the PDF file
        output_dir: Optional directory to save extracted images
    
    Returns:
        List of dictionaries with image info and paths
    
    Raises:
        ImportError: If PyMuPDF is not available
        Exception: If PDF cannot be read or is corrupted
    """
    if not PYMUPDF_AVAILABLE:
        raise ImportError("PyMuPDF is not installed. Install with: pip install PyMuPDF")
    
    from pathlib import Path as PathLib
    import hashlib
    
    extracted_images = []
    
    if output_dir is None:
        # Use temp_media directory
        from app.config import settings
        output_dir = settings.MEDIA_TEMP_DIR / f"pdf_images_{PathLib(file_path).stem}"
    
    output_path = PathLib(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                # Get list of images on the page
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]  # Image XREF number
                    
                    # Extract image
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Calculate hash for unique filename
                    img_hash = hashlib.md5(image_bytes).hexdigest()[:12]
                    
                    # Save image
                    image_filename = f"page{page_num}_img{img_index}_{img_hash}.{image_ext}"
                    image_path = output_path / image_filename
                    
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    extracted_images.append({
                        "page": page_num,
                        "index": img_index,
                        "path": str(image_path),
                        "format": image_ext,
                        "size_bytes": len(image_bytes),
                        "xref": xref
                    })
                    
                    logger.debug(
                        f"Extracted image from PDF page {page_num}: "
                        f"{image_filename} ({len(image_bytes)} bytes)"
                    )
    
    except Exception as e:
        logger.error(f"Failed to extract images from PDF: {e}")
        raise
    
    logger.info(f"Extracted {len(extracted_images)} images from PDF: {file_path}")
    return extracted_images


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Extracts text from paragraphs and tables.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        Extracted text content
    
    Raises:
        ImportError: If python-docx is not available
        Exception: If DOCX cannot be read or is corrupted
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    
    text_parts = []
    
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        logger.debug(f"Extracted {len(text_parts)} text blocks from DOCX")
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise
    
    return '\n'.join(text_parts)


def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
    """
    Split text into overlapping chunks.
    
    Uses sentence boundaries when possible to avoid mid-sentence breaks.
    
    Args:
        text: Text to chunk
        chunk_size: Maximum characters per chunk (default from settings)
        overlap: Number of overlapping characters (default from settings)
    
    Returns:
        List of text chunks
    
    Example:
        >>> text = "Hello world. This is a test. More text here."
        >>> chunks = chunk_text(text, chunk_size=20, overlap=5)
        >>> len(chunks) >= 2
        True
    """
    if chunk_size is None:
        chunk_size = settings.CHUNK_SIZE
    if overlap is None:
        overlap = settings.CHUNK_OVERLAP
    
    if not text or len(text) <= chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # If this isn't the last chunk, try to break at sentence boundary
        if end < len(text):
            # Look for sentence-ending punctuation in the last 50 chars
            search_start = max(start, end - 50)
            search_text = text[search_start:end]
            
            # Find last sentence boundary
            for punct in ['. ', '! ', '? ', '\n\n', '\n']:
                idx = search_text.rfind(punct)
                if idx != -1:
                    end = search_start + idx + len(punct)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = end - overlap if end < len(text) else end
    
    logger.debug(f"Split text into {len(chunks)} chunks")
    return chunks


def extract_file_text(
    file_path: str,
    chunk_size: Optional[int] = None,
    overlap: Optional[int] = None
) -> tuple[List[FileChunk], FileInfo]:
    """
    Extract text from a file and chunk it.
    
    Supports TXT, MD, PDF, and DOCX files.
    
    Args:
        file_path: Path to the file
        chunk_size: Maximum characters per chunk (default from settings)
        overlap: Number of overlapping characters (default from settings)
    
    Returns:
        Tuple of (list of FileChunk objects, FileInfo metadata)
    
    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
        Exception: For extraction errors
    
    Example:
        >>> chunks, info = extract_file_text("example.txt")
        >>> info.extraction_success
        True
        >>> len(chunks) > 0
        True
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    file_ext = Path(file_path).suffix.lower().lstrip('.')
    
    # Validate file type
    if f".{file_ext}" not in settings.ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    # Calculate file hash
    try:
        file_hash = hash_file_sha256(file_path)
        logger.debug(f"File hash: {file_hash}")
    except Exception as e:
        logger.error(f"Failed to hash file: {e}")
        raise
    
    # Extract text based on file type
    extraction_error = None
    text = ""
    
    try:
        if file_ext == 'txt':
            text = extract_text_from_txt(file_path)
        elif file_ext == 'md':
            text = extract_text_from_md(file_path)
        elif file_ext == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif file_ext == 'docx':
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
        
        extraction_success = True
        logger.info(f"Extracted {len(text)} characters from {file_ext} file")
        
    except Exception as e:
        extraction_success = False
        extraction_error = str(e)
        logger.error(f"Extraction failed: {e}")
    
    # Chunk the text
    if text:
        text_chunks = chunk_text(text, chunk_size, overlap)
    else:
        text_chunks = []
    
    # Create FileChunk objects
    file_chunks = [
        FileChunk(
            content=chunk,
            source=file_path,
            hash=file_hash,
            chunk_id=i
        )
        for i, chunk in enumerate(text_chunks)
    ]
    
    # Create FileInfo metadata
    file_info = FileInfo(
        original_path=file_path,
        hash=file_hash,
        type=file_ext,
        chunk_count=len(file_chunks),
        extraction_success=extraction_success,
        extraction_error=extraction_error
    )
    
    return file_chunks, file_info


def validate_file(file_path: str) -> tuple[bool, Optional[str]]:
    """
    Validate a file before extraction.
    
    Checks file existence, size, and extension.
    
    Args:
        file_path: Path to the file to validate
    
    Returns:
        Tuple of (is_valid, error_message)
    
    Example:
        >>> valid, error = validate_file("example.txt")
        >>> valid
        True
        >>> error is None
        True
    """
    # Check existence
    if not os.path.exists(file_path):
        return False, f"File not found: {file_path}"
    
    if not os.path.isfile(file_path):
        return False, f"Not a file: {file_path}"
    
    # Check size
    file_size = os.path.getsize(file_path)
    if file_size > settings.MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        return False, f"File too large: {size_mb:.2f}MB (max: {settings.MAX_FILE_SIZE_MB}MB)"
    
    # Check extension
    if not settings.is_allowed_extension(file_path):
        ext = Path(file_path).suffix
        return False, f"Unsupported file type: {ext}"
    
    # Check library availability for specific types
    file_ext = Path(file_path).suffix.lower().lstrip('.')
    if file_ext == 'pdf' and not PYMUPDF_AVAILABLE:
        return False, "PDF extraction not available (PyMuPDF not installed)"
    if file_ext == 'docx' and not PYTHON_DOCX_AVAILABLE:
        return False, "DOCX extraction not available (python-docx not installed)"
    
    return True, None

